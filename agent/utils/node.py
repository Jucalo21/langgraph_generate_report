from agent.utils.tools import definir_llm, tavily_search
from docx import Document as DocxDocument
from dotenv import load_dotenv
from agent.utils.state import ResearchState, Document, Query
import os, json

load_dotenv()


def create_queries(state: ResearchState):
    client = definir_llm()

    prompt = f"""
    Crea un conjunto de consultas que ayuden a investigar en profundidad el tema especificado.

    Utiliza la siguiente información para generar las consultas:

    Temática: <theme>{state.researcher.theme}</theme>
    Número de consultas: <number_queries>{state.researcher.number_queries}</number_queries>
    Cada consulta debe ser clara, precisa y enfocada en aspectos relevantes del tema. 
    Asegúrate de que las preguntas estén orientadas a extraer información detallada y fomentar un análisis profundo. No incluyas explicaciones, comentarios u otro contenido adicional; solo entrega las consultas.

    Steps
    Revisa la temática proporcionada.
    Identifica las áreas clave del tema que requieren investigación.
    Formula <number_queries>{state.researcher.number_queries}</number_queries> preguntas de investigación basadas en la temática.
    Verifica que cada pregunta esté diseñada para profundizar en el análisis del tema.
    Confirma que la salida incluya únicamente las consultas, sin encabezados ni contenido extra.
    Output Format
    La salida debe consistir en una lista de consultas, presentadas de forma secuencial (por ejemplo, numeradas o separadas por guiones). No se deben incluir encabezados, explicaciones o comentarios adicionales.

    Ejemplos
    Ejemplo de entrada:
    Temática: "Cambio Climático"
    Número de consultas: 3

    Ejemplo de salida:

    ¿Cuáles son las principales causas del cambio climático y cómo afectan a distintas regiones?
    ¿Qué estrategias se están implementando globalmente para mitigar los efectos del cambio climático?
    ¿Cómo se relacionan las actividades humanas con el aumento de la temperatura global?
    Notas
    Utiliza los marcadores de estado proporcionados (<theme> y <number_queries>) para incorporar la información correspondiente.
    Asegúrate de entregar únicamente las consultas solicitadas, sin agregar contenido adicional.
        """

    response = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=prompt,
        config={
            "response_mime_type": "application/json",
            "response_schema": Query,
        },
    )

    parsed = json.loads(response.text)
    state.researcher.queries = Query(queries=parsed.get("queries", []))

    return state


def investigate_queries(state: ResearchState):
    query_list = state.researcher.queries.queries

    info_acumulada = ""
    for query in query_list:
        result = tavily_search(query)
        results_list = result.get("results", [])

        for item in results_list:
            info_acumulada += "\n".join(
                [
                    f"Título: {item.get('title', '')}\n",
                    f"URL: {item.get('url', '')}\n",
                    f"Contenido: {item.get('raw_content' if item.get('raw_content') else item.get('content'))}\n",
                ]
            )

    state.researcher.info_documento = info_acumulada.strip()
    return state


def create_report(state: ResearchState) -> Document:

    # Abre el archivo en modo escritura ('w')
    with open("mi_archivo.txt", "w", encoding="utf-8") as archivo:
        archivo.write(state.researcher.info_documento)  # Escribe una línea

    prompt = f"""Crear un reporte de investigación bien estructurado, basado en la temática e información acumulada proporcionada.
    Utiliza la información dada para desarrollar un reporte que se divida en las siguientes secciones, 
    donde cada sección cumple un objetivo específico y deben tener un mínimo de 1000 palabras y maximo de 2000:

    Título: Debe ser atractivo, descriptivo y reflejar la temática de la investigación.
    Introducción: Presenta el tema, el contexto y la relevancia del estudio.
    Cuerpo: Desarrolla un análisis detallado y presenta los datos acumulados. Asegúrate de que el proceso de razonamiento y análisis se exponga de forma clara y completa antes de llegar a conclusiones.
    Conclusión: Resume los resultados obtenidos y destaca su relevancia. Esta sección debe aparecer al final, después de haber desarrollado todo el análisis.
    Utiliza la siguiente información como base:

    Temática: 
    <theme>
    {state.researcher.theme}
    </theme>
    
    Información acumulada: 
    <info_documento>
    {state.researcher.info_documento}
    </info_documento>
    
    Steps:
    Análisis inicial: Revisa y comprende la temática y la información acumulada.
    Título: Crea un título atractivo y descriptivo que capte la esencia de la investigación.
    Introducción: Escribe una introducción que contextualice el tema y explique su relevancia.
    Cuerpo: Desarrolla el cuerpo del reporte presentando los datos y realizando un análisis exhaustivo. Asegúrate de que todo el proceso de razonamiento se presente de forma clara y detallada antes de llegar a cualquier conclusión.
    Conclusión: Finaliza con una conclusión que resuma los hallazgos y destaque la importancia de los resultados obtenidos.
    Revisión: Verifica que cada sección cumpla con el mínimo de 500 palabras y que la estructura del reporte sea lógica y coherente.
    
    Output Format:
    El reporte final debe estar estructurado en secciones con los siguientes encabezados:
    Título
    Introducción
    Cuerpo
    Conclusión

    Cada sección debe presentarse en texto plano, sin bloques de código, y contener al menos 500 palabras.

    Ejemplos
    Ejemplo de estructura:

    Título: "[Título descriptivo y atractivo basado en la temática]"
    Introducción: "[Introducción detallada que presenta el tema y su relevancia, incluyendo contexto y objetivos]"
    Cuerpo: "[Desarrollo extenso que expone los datos, análisis y razonamiento de la investigación, presentando el proceso analítico antes de llegar a conclusiones]"
    Conclusión: "[Resumen final que consolida los resultados y explica la importancia de los hallazgos]"
    (Los ejemplos reales deben ser más extensos, asegurándose de que cada sección cumpla con el mínimo de 500 palabras).

    Notas:
    Asegúrate de que el proceso de razonamiento y análisis se presente antes de exponer cualquier conclusión.
    Revisa que la estructura, la claridad y el contenido sean coherentes y completos, cumpliendo con los requisitos establecidos.
    """

    client = definir_llm()

    response = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=prompt,
        config={
            "response_mime_type": "application/json",
            "response_schema": Document,
        },
    )

    parsed = json.loads(response.text)
    state.document = Document(**parsed)
    return state


def create_document(state: ResearchState) -> None:

    # Crear un nuevo documento de word
    doc = DocxDocument()

    # Agregar el titulo
    doc.add_heading(state.document.title, level=1)

    # Agregar la introduccion
    doc.add_paragraph(state.document.introduction)

    # Agregar el cuerpo
    doc.add_paragraph(state.document.body)

    # Agregar la conclusion
    doc.add_paragraph(state.document.conclusion)

    # Guardar el documento
    doc.save("output/reporte.docx")
    print("Documento Generado")
    return state
